import win32com.client
import re
from docx import Document
from flask import Flask, render_template, request, send_file
import os

# Initialize Flask app
app = Flask(__name__)

def extract_vba_code(excel_file):
    try:
        excel = win32com.client.Dispatch("Excel.Application")
        excel.Visible = False
        workbook = excel.Workbooks.Open(excel_file)
        vba_project = workbook.VBProject
        
        vba_code = {}
        for component in vba_project.VBComponents:
            if component.Type in [1, 2, 3]:  # Standard, Class, and Module
                vba_code[component.Name] = component.CodeModule.Lines(1, component.CodeModule.CountOfLines)
        
        workbook.Close(False)
        excel.Quit()
        return vba_code
    except Exception as e:
        print(f"Error extracting VBA code: {e}")
        return None

def analyze_vba_code(vba_code):
    analysis = {}
    for module, code in vba_code.items():
        procedures = re.findall(r"Sub\s+(\w+)", code)
        variables = re.findall(r"Dim\s+(\w+)", code)
        analysis[module] = {
            'Procedures': procedures,
            'Variables': variables
        }
    return analysis

def generate_documentation(vba_code, analysis, output_file):
    document = Document()
    document.add_heading('VBA Macro Documentation', 0)

    for module, code in vba_code.items():
        document.add_heading(f'Module: {module}', level=1)
        document.add_paragraph('Code:')
        document.add_paragraph(code)
        
        if module in analysis:
            document.add_heading('Analysis', level=2)
            procedures = analysis[module].get('Procedures', [])
            variables = analysis[module].get('Variables', [])
            
            document.add_heading('Procedures:', level=3)
            for proc in procedures:
                document.add_paragraph(proc)
            
            document.add_heading('Variables:', level=3)
            for var in variables:
                document.add_paragraph(var)

    document.save(output_file)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return 'No file part'
    file = request.files['file']
    if file.filename == '':
        return 'No selected file'
    if file:
        excel_file = os.path.join('uploads', file.filename)
        os.makedirs('uploads', exist_ok=True)
        file.save(excel_file)
        output_file = os.path.join('outputs', f'Documentation_{os.path.splitext(file.filename)[0]}.docx')
        os.makedirs('outputs', exist_ok=True)
        analysis = analyze_vba_code(extract_vba_code(excel_file))
        generate_documentation(extract_vba_code(excel_file), analysis, output_file)
        return send_file(output_file, as_attachment=True)

if __name__ == "__main__":
    app.run(debug=True)
