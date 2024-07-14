import os
import win32com.client
import pythoncom  # Import pythoncom for CoInitialize
import pandas as pd
from docx import Document
import matplotlib.pyplot as plt
import seaborn as sns
from flask import Flask, render_template, request, send_file

# Initialize COM
try:
    pythoncom.CoInitialize()
except pythoncom.com_error as ce:
    print(f"Failed to initialize COM: {ce}")

# Flask app setup
app = Flask(__name__)

# Function to extract VBA code from Excel
def extract_vba_code(excel_file):
    excel = win32com.client.Dispatch("Excel.Application")
    excel.Visible = False
    workbook = None
    
    try:
        workbook = excel.Workbooks.Open(excel_file)
        vba_project = workbook.VBProject

        vba_code = {}
        for component in vba_project.VBComponents:
            if component.Type in [1, 2, 3]:  # Standard, Class, and Module
                vba_code[component.Name] = component.CodeModule.Lines(1, component.CodeModule.CountOfLines)
    except Exception as e:
        print(f"Error extracting VBA code: {e}")
    finally:
        if workbook:
            workbook.Close(SaveChanges=False)
        excel.Quit()
        
    return vba_code

# Function to parse VBA code structure
def parse_vba_code(vba_code):
    # Define your VBA code parsing logic here using pyparsing or other suitable libraries
    # Example parsing logic goes here
    parsed_code = {}
    for module_name, code in vba_code.items():
        # Implement your parsing logic for each module
        parsed_code[module_name] = code  # Placeholder for actual parsing
    
    return parsed_code

# Function to extract data from Excel
def extract_data(excel_file):
    try:
        df = pd.read_excel(excel_file, engine='openpyxl')  # You can use 'xlrd' or 'openpyxl' depending on your Excel file format
        return df
    except Exception as e:
        print(f"Error reading Excel file: {e}")
        return None

# Function to generate Word documentation
def generate_documentation(vba_code, df, output_file):
    document = Document()
    document.add_heading('Excel VBA Macro and Dataset Documentation', 0)

    document.add_heading('VBA Macros Overview', level=1)
    for module_name, code in vba_code.items():
        document.add_heading(f'Module: {module_name}', level=2)
        document.add_paragraph('Code:')
        document.add_paragraph(f'{code}')  # Example, replace with actual parsed code
    
    document.add_heading('Dataset Overview', level=1)
    if df is not None:
        document.add_paragraph(f'The dataset contains {len(df)} records and {len(df.columns)} columns.')
        document.add_heading('Columns:', level=2)
        for col in df.columns:
            document.add_paragraph(col)
        
        document.add_heading('Sample Records:', level=2)
        for index, row in df.head().iterrows():
            document.add_paragraph(str(row.to_dict()))
    else:
        document.add_paragraph('Error: Dataset could not be loaded.')

    document.save(output_file)

# Function to create visualizations
def create_visualizations(df, output_dir):
    if df is None:
        return
    
    os.makedirs(output_dir, exist_ok=True)
    
    plt.figure(figsize=(10, 6))
    sns.countplot(data=df, x='Genre')
    plt.title('Count of Books by Genre')
    plt.xticks(rotation=45)
    plt.tight_layout()
    plt.savefig(os.path.join(output_dir, 'genre_distribution.png'))

    plt.figure(figsize=(10, 6))
    sns.scatterplot(data=df, x='Year', y='User Rating', hue='Genre')
    plt.title('User Rating by Year and Genre')
    plt.tight_layout()
    plt.savefig(os.path.join(output_dir, 'rating_by_year_genre.png'))

    plt.figure(figsize=(10, 6))
    sns.boxplot(data=df, x='Genre', y='Price')
    plt.title('Price Distribution by Genre')
    plt.xticks(rotation=45)
    plt.tight_layout()
    plt.savefig(os.path.join(output_dir, 'price_distribution_by_genre.png'))

# Main function to run the entire process
def main(excel_file, output_file, output_dir):
    vba_code = extract_vba_code(excel_file)
    parsed_code = parse_vba_code(vba_code)
    df = extract_data(excel_file)
    generate_documentation(parsed_code, df, output_file)
    create_visualizations(df, output_dir)

# Flask routes
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return 'No file part'
    file = request.files['file']
    if file.filename == '':
        return 'No selected file'
    if file:
        excel_file = os.path.join('uploads', file.filename)
        file.save(excel_file)
        output_file = os.path.join('outputs', f'Documentation_{os.path.splitext(file.filename)[0]}.docx')
        output_dir = os.path.join('outputs', os.path.splitext(file.filename)[0])
        os.makedirs(output_dir, exist_ok=True)
        main(excel_file, output_file, output_dir)
        return send_file(output_file, as_attachment=True)

if __name__ == "__main__":
    app.run(debug=True)
